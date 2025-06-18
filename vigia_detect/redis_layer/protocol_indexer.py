"""
Enhanced Protocol Indexer - Indexador de protocolos médicos con capacidades avanzadas
Implementación completa para indexación y búsqueda de protocolos médicos en PDFs y documentos.

Características:
- Extracción de texto de PDFs médicos
- Indexación con metadatos clínicos
- Búsqueda por tipo de protocolo
- Cache inteligente de resultados
- Soporte para múltiples formatos
"""

import asyncio
import json
import hashlib
import re
from datetime import datetime, timezone, timedelta
from typing import Dict, Any, Optional, List, Tuple, Union
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
import logging

# Document processing
try:
    import fitz  # PyMuPDF for PDF processing
except ImportError:
    # Mock para desarrollo sin PyMuPDF
    class MockFitz:
        def open(self, *args, **kwargs):
            return MockDocument()
    
    class MockDocument:
        def load_page(self, page_num):
            return MockPage()
        
        def close(self):
            pass
    
    class MockPage:
        def get_text(self):
            return "Mock medical protocol text for testing"
    
    fitz = MockFitz()
from docx import Document  # python-docx for Word documents

# Redis imports
import redis.asyncio as redis
from redis.exceptions import RedisError

from ..utils.secure_logger import SecureLogger
from ..utils.error_handling import handle_exceptions

logger = SecureLogger("protocol_indexer_enhanced")


class ProtocolType(Enum):
    """Tipos de protocolos médicos."""
    PREVENTION = "prevention"
    TREATMENT = "treatment"
    DIAGNOSTIC = "diagnostic"
    MEDICATION = "medication"
    SURGERY = "surgery"
    EMERGENCY = "emergency"
    NURSING_CARE = "nursing_care"
    INFECTION_CONTROL = "infection_control"
    QUALITY_ASSURANCE = "quality_assurance"
    GENERAL = "general"


class DocumentFormat(Enum):
    """Formatos de documento soportados."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    MD = "md"


@dataclass
class ProtocolIndexConfig:
    """Configuración para indexación de protocolos."""
    cache_ttl: int = 7200  # 2 horas
    max_extract_length: int = 50000  # Máximo caracteres por documento
    chunk_size: int = 1000  # Tamaño de chunks para indexación
    overlap_size: int = 100  # Solapamiento entre chunks
    min_chunk_size: int = 200  # Tamaño mínimo de chunk
    supported_formats: List[str] = None
    
    def __post_init__(self):
        if self.supported_formats is None:
            self.supported_formats = ["pdf", "docx", "txt", "html", "md"]


@dataclass
class ProtocolDocument:
    """Documento de protocolo médico."""
    doc_id: str
    title: str
    content: str
    protocol_type: ProtocolType
    metadata: Dict[str, Any]
    chunks: List[str]
    language: str = "es"
    file_path: Optional[str] = None
    created_at: datetime = None
    last_updated: datetime = None
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now(timezone.utc)
        if self.last_updated is None:
            self.last_updated = self.created_at


@dataclass
class ProtocolSearchResult:
    """Resultado de búsqueda de protocolo."""
    doc_id: str
    title: str
    content_snippet: str
    protocol_type: ProtocolType
    metadata: Dict[str, Any]
    relevance_score: float
    matched_chunks: List[str]
    language: str


class EnhancedProtocolIndexer:
    """Indexador enhanced de protocolos médicos."""
    
    def __init__(self, config: Optional[ProtocolIndexConfig] = None):
        """
        Inicializar indexador enhanced.
        
        Args:
            config: Configuración del indexador
        """
        self.config = config or ProtocolIndexConfig()
        self.redis_client: Optional[redis.Redis] = None
        self.protocol_store: Dict[str, ProtocolDocument] = {}
        
        # Índices invertidos para búsqueda rápida
        self.term_index: Dict[str, set] = {}  # término -> set de doc_ids
        self.type_index: Dict[ProtocolType, set] = {}  # tipo -> set de doc_ids
        
        # Cache de búsquedas
        self.search_cache: Dict[str, List[Dict[str, Any]]] = {}
        
        # Estadísticas
        self.stats = {
            "documents_indexed": 0,
            "total_chunks": 0,
            "search_queries": 0,
            "cache_hits": 0,
            "extraction_errors": 0
        }
        
        logger.audit("enhanced_protocol_indexer_initialized", {
            "supported_formats": self.config.supported_formats,
            "chunk_size": self.config.chunk_size
        })
    
    async def initialize(self):
        """Inicializar servicios."""
        try:
            # Inicializar Redis
            await self._initialize_redis()
            
            # Cargar índices existentes
            await self._load_existing_indices()
            
            logger.info("Enhanced protocol indexer initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize enhanced protocol indexer: {e}")
            raise
    
    async def _initialize_redis(self):
        """Inicializar conexión a Redis."""
        try:
            self.redis_client = redis.Redis(
                host='localhost',
                port=6379,
                decode_responses=True,
                socket_timeout=5.0,
                socket_connect_timeout=5.0,
                retry_on_timeout=True
            )
            
            # Probar conexión
            await self.redis_client.ping()
            logger.info("Redis connection established for protocol indexer")
            
        except Exception as e:
            logger.warning(f"Redis connection failed for protocol indexer: {e}")
            self.redis_client = None
    
    async def _load_existing_indices(self):
        """Cargar índices existentes desde Redis."""
        if not self.redis_client:
            logger.warning("Redis not available, skipping index loading")
            return
        
        try:
            # Cargar documentos
            doc_keys = await self.redis_client.keys("protocol_doc:*")
            
            for key in doc_keys:
                try:
                    doc_data = await self.redis_client.hgetall(key)
                    if doc_data:
                        doc = self._deserialize_protocol_document(doc_data)
                        if doc:
                            self.protocol_store[doc.doc_id] = doc
                            self._update_indices(doc)
                            self.stats["documents_indexed"] += 1
                            self.stats["total_chunks"] += len(doc.chunks)
                
                except Exception as e:
                    logger.warning(f"Failed to load protocol document {key}: {e}")
                    continue
            
            logger.info(f"Loaded {self.stats['documents_indexed']} protocol documents")
            
        except Exception as e:
            logger.error(f"Failed to load existing indices: {e}")
    
    @handle_exceptions(logger)
    async def index_document_file(self, 
                                file_path: Union[str, Path],
                                protocol_type: ProtocolType,
                                metadata: Optional[Dict[str, Any]] = None,
                                title: Optional[str] = None) -> str:
        """
        Indexar archivo de protocolo médico.
        
        Args:
            file_path: Ruta al archivo
            protocol_type: Tipo de protocolo
            metadata: Metadatos adicionales
            title: Título del documento (opcional)
            
        Returns:
            ID del documento indexado
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Detectar formato
            file_format = file_path.suffix.lower()[1:]  # Remover el punto
            
            if file_format not in self.config.supported_formats:
                raise ValueError(f"Unsupported file format: {file_format}")
            
            # Extraer contenido
            content = await self._extract_content(file_path, file_format)
            
            # Usar nombre de archivo como título si no se proporciona
            if not title:
                title = file_path.stem
            
            # Metadatos por defecto
            file_metadata = {
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "file_format": file_format,
                "extraction_date": datetime.now(timezone.utc).isoformat(),
                **(metadata or {})
            }
            
            # Indexar contenido
            doc_id = await self.index_document_content(
                content=content,
                title=title,
                protocol_type=protocol_type,
                metadata=file_metadata,
                file_path=str(file_path)
            )
            
            logger.audit("protocol_file_indexed", {
                "doc_id": doc_id,
                "file_path": str(file_path),
                "file_format": file_format,
                "protocol_type": protocol_type.value,
                "content_length": len(content)
            })
            
            return doc_id
            
        except Exception as e:
            self.stats["extraction_errors"] += 1
            logger.error(f"Failed to index document file {file_path}: {e}")
            raise
    
    @handle_exceptions(logger)
    async def index_document_content(self,
                                   content: str,
                                   title: str,
                                   protocol_type: ProtocolType,
                                   metadata: Optional[Dict[str, Any]] = None,
                                   file_path: Optional[str] = None,
                                   language: str = "es") -> str:
        """
        Indexar contenido de protocolo médico.
        
        Args:
            content: Contenido del documento
            title: Título del documento
            protocol_type: Tipo de protocolo
            metadata: Metadatos adicionales
            file_path: Ruta del archivo (opcional)
            language: Idioma del documento
            
        Returns:
            ID del documento indexado
        """
        try:
            # Generar ID único
            doc_id = hashlib.sha256(
                f"{title}_{content[:100]}_{datetime.now().isoformat()}".encode()
            ).hexdigest()[:16]
            
            # Limpiar y procesar contenido
            cleaned_content = self._clean_medical_content(content)
            
            # Dividir en chunks
            chunks = self._create_chunks(cleaned_content)
            
            # Crear documento
            doc = ProtocolDocument(
                doc_id=doc_id,
                title=title,
                content=cleaned_content,
                protocol_type=protocol_type,
                metadata=metadata or {},
                chunks=chunks,
                language=language,
                file_path=file_path
            )
            
            # Almacenar documento
            self.protocol_store[doc_id] = doc
            
            # Actualizar índices
            self._update_indices(doc)
            
            # Persistir en Redis
            await self._persist_protocol_document(doc)
            
            # Actualizar estadísticas
            self.stats["documents_indexed"] += 1
            self.stats["total_chunks"] += len(chunks)
            
            logger.audit("protocol_content_indexed", {
                "doc_id": doc_id,
                "title": title,
                "protocol_type": protocol_type.value,
                "chunks_created": len(chunks),
                "content_length": len(cleaned_content)
            })
            
            return doc_id
            
        except Exception as e:
            logger.error(f"Failed to index document content: {e}")
            raise
    
    async def _extract_content(self, file_path: Path, file_format: str) -> str:
        """Extraer contenido de archivo según su formato."""
        try:
            if file_format == "pdf":
                return await self._extract_pdf_content(file_path)
            elif file_format == "docx":
                return await self._extract_docx_content(file_path)
            elif file_format in ["txt", "md"]:
                return await self._extract_text_content(file_path)
            elif file_format == "html":
                return await self._extract_html_content(file_path)
            else:
                raise ValueError(f"Unsupported format: {file_format}")
                
        except Exception as e:
            logger.error(f"Content extraction failed for {file_path}: {e}")
            raise
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """Extraer contenido de archivo PDF."""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_pdf():
                doc = fitz.open(str(file_path))
                content = ""
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    content += page.get_text()
                    
                    # Limitar longitud
                    if len(content) > self.config.max_extract_length:
                        content = content[:self.config.max_extract_length]
                        break
                
                doc.close()
                return content
            
            content = await loop.run_in_executor(None, extract_pdf)
            return content.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """Extraer contenido de archivo DOCX."""
        try:
            loop = asyncio.get_event_loop()
            
            def extract_docx():
                doc = Document(str(file_path))
                content = ""
                
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                    
                    # Limitar longitud
                    if len(content) > self.config.max_extract_length:
                        content = content[:self.config.max_extract_length]
                        break
                
                return content
            
            content = await loop.run_in_executor(None, extract_docx)
            return content.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    async def _extract_text_content(self, file_path: Path) -> str:
        """Extraer contenido de archivo de texto."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read(self.config.max_extract_length)
            return content.strip()
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            raise
    
    async def _extract_html_content(self, file_path: Path) -> str:
        """Extraer contenido de archivo HTML."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remover scripts y estilos
            for script in soup(["script", "style"]):
                script.extract()
            
            content = soup.get_text()
            
            # Limpiar espacios en blanco excesivos
            content = re.sub(r'\s+', ' ', content).strip()
            
            if len(content) > self.config.max_extract_length:
                content = content[:self.config.max_extract_length]
            
            return content
            
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            raise
    
    def _clean_medical_content(self, content: str) -> str:
        """Limpiar contenido médico para indexación."""
        # Remover caracteres de control
        content = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', content)
        
        # Normalizar espacios en blanco
        content = re.sub(r'\s+', ' ', content)
        
        # Preservar términos médicos importantes
        medical_terms = [
            r'\b(LPP|pressure injury|lesión por presión)\b',
            r'\b(grado|grade|estadio|stage)\s+[IVX1-4]+\b',
            r'\b(protocolo|protocol|procedimiento|procedure)\b',
            r'\b(tratamiento|treatment|terapia|therapy)\b'
        ]
        
        # Marcar términos médicos importantes para preservarlos
        for i, pattern in enumerate(medical_terms):
            content = re.sub(pattern, f'MEDTERM{i}_\\g<0>_MEDTERM{i}', content, flags=re.IGNORECASE)
        
        # Limpiar caracteres especiales pero preservar marcadores
        content = re.sub(r'[^\w\s\-.,;:()\[\]MEDTERM_]', '', content)
        
        # Restaurar términos médicos
        for i in range(len(medical_terms)):
            content = re.sub(f'MEDTERM{i}_(.+?)_MEDTERM{i}', r'\1', content)
        
        return content.strip()
    
    def _create_chunks(self, content: str) -> List[str]:
        """Dividir contenido en chunks para indexación."""
        chunks = []
        words = content.split()
        
        if not words:
            return chunks
        
        # Calcular palabras por chunk aproximadamente
        words_per_chunk = self.config.chunk_size // 5  # Estimación: 5 caracteres por palabra
        overlap_words = self.config.overlap_size // 5
        
        start = 0
        while start < len(words):
            end = min(start + words_per_chunk, len(words))
            chunk_words = words[start:end]
            chunk_text = ' '.join(chunk_words)
            
            # Verificar tamaño mínimo
            if len(chunk_text) >= self.config.min_chunk_size:
                chunks.append(chunk_text)
            
            # Calcular siguiente inicio con solapamiento
            if end >= len(words):
                break
            
            start = end - overlap_words
            if start <= 0:
                start = end
        
        return chunks
    
    def _update_indices(self, doc: ProtocolDocument):
        """Actualizar índices invertidos."""
        # Índice por tipo
        if doc.protocol_type not in self.type_index:
            self.type_index[doc.protocol_type] = set()
        self.type_index[doc.protocol_type].add(doc.doc_id)
        
        # Índice de términos
        all_text = f"{doc.title} {doc.content}".lower()
        terms = self._extract_search_terms(all_text)
        
        for term in terms:
            if term not in self.term_index:
                self.term_index[term] = set()
            self.term_index[term].add(doc.doc_id)
    
    def _extract_search_terms(self, text: str) -> set:
        """Extraer términos de búsqueda del texto."""
        # Tokenizar y limpiar
        words = re.findall(r'\b\w{3,}\b', text.lower())  # Palabras de 3+ caracteres
        
        # Filtrar stop words comunes
        stop_words = {
            'and', 'the', 'for', 'are', 'with', 'this', 'that', 'from', 'they', 'been',
            'have', 'has', 'had', 'will', 'would', 'could', 'should', 'may', 'might',
            'una', 'los', 'las', 'del', 'por', 'para', 'con', 'que', 'como', 'este',
            'esta', 'estos', 'estas', 'pero', 'desde', 'hasta', 'durante', 'entre'
        }
        
        # Incluir términos médicos importantes
        medical_terms = {
            'lpp', 'lesion', 'presion', 'pressure', 'injury', 'ulcer', 'wound',
            'tratamiento', 'treatment', 'protocolo', 'protocol', 'procedimiento',
            'medicamento', 'medication', 'antibiotico', 'antibiotic', 'dolor', 'pain'
        }
        
        filtered_terms = set()
        for word in words:
            if word not in stop_words and (len(word) >= 4 or word in medical_terms):
                filtered_terms.add(word)
        
        return filtered_terms
    
    @handle_exceptions(logger)
    async def search_protocols(self,
                             query: str,
                             protocol_type: Optional[str] = None,
                             query_type: Optional[str] = None,
                             max_results: int = 10,
                             language: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        Buscar protocolos médicos.
        
        Args:
            query: Consulta de búsqueda
            protocol_type: Filtrar por tipo de protocolo
            query_type: Tipo de consulta para contexto
            max_results: Número máximo de resultados
            language: Filtrar por idioma
            
        Returns:
            Lista de protocolos encontrados
        """
        try:
            self.stats["search_queries"] += 1
            
            # Verificar cache
            cache_key = self._get_search_cache_key(query, protocol_type, query_type, max_results, language)
            
            if cache_key in self.search_cache:
                self.stats["cache_hits"] += 1
                return self.search_cache[cache_key]
            
            # Realizar búsqueda
            results = await self._execute_search(
                query=query,
                protocol_type=protocol_type,
                query_type=query_type,
                max_results=max_results,
                language=language
            )
            
            # Serializar resultados para compatibilidad
            serialized_results = [self._serialize_search_result(r) for r in results]
            
            # Guardar en cache
            self.search_cache[cache_key] = serialized_results
            
            # Limpiar cache si es muy grande
            if len(self.search_cache) > 100:
                # Remover entradas más antiguas (implementación simple)
                oldest_keys = list(self.search_cache.keys())[:20]
                for key in oldest_keys:
                    del self.search_cache[key]
            
            logger.audit("protocol_search_completed", {
                "query_length": len(query),
                "results_found": len(results),
                "protocol_type_filter": protocol_type,
                "cache_hit": False
            })
            
            return serialized_results
            
        except Exception as e:
            logger.error(f"Protocol search failed: {e}")
            return []
    
    async def _execute_search(self,
                            query: str,
                            protocol_type: Optional[str],
                            query_type: Optional[str],
                            max_results: int,
                            language: Optional[str]) -> List[ProtocolSearchResult]:
        """Ejecutar búsqueda de protocolos."""
        try:
            # Extraer términos de búsqueda
            search_terms = self._extract_search_terms(query.lower())
            
            if not search_terms:
                return []
            
            # Encontrar documentos candidatos
            candidate_docs = self._find_candidate_documents(search_terms, protocol_type)
            
            if not candidate_docs:
                return []
            
            # Calcular relevancia para cada documento
            scored_results = []
            
            for doc_id in candidate_docs:
                if doc_id not in self.protocol_store:
                    continue
                
                doc = self.protocol_store[doc_id]
                
                # Filtrar por idioma si se especifica
                if language and doc.language != language:
                    continue
                
                # Calcular score de relevancia
                relevance_score = self._calculate_relevance_score(query, search_terms, doc)
                
                if relevance_score > 0.1:  # Umbral mínimo
                    # Encontrar chunks relevantes
                    matched_chunks = self._find_matching_chunks(search_terms, doc.chunks)
                    
                    # Crear snippet
                    snippet = self._create_content_snippet(query, doc.content)
                    
                    result = ProtocolSearchResult(
                        doc_id=doc.doc_id,
                        title=doc.title,
                        content_snippet=snippet,
                        protocol_type=doc.protocol_type,
                        metadata=doc.metadata,
                        relevance_score=relevance_score,
                        matched_chunks=matched_chunks[:3],  # Top 3 chunks
                        language=doc.language
                    )
                    
                    scored_results.append(result)
            
            # Ordenar por relevancia y limitar resultados
            scored_results.sort(key=lambda x: x.relevance_score, reverse=True)
            return scored_results[:max_results]
            
        except Exception as e:
            logger.error(f"Search execution failed: {e}")
            return []
    
    def _find_candidate_documents(self, search_terms: set, protocol_type: Optional[str]) -> set:
        """Encontrar documentos candidatos basados en términos de búsqueda."""
        candidate_docs = set()
        
        # Buscar por términos
        for term in search_terms:
            if term in self.term_index:
                candidate_docs.update(self.term_index[term])
        
        # Filtrar por tipo de protocolo si se especifica
        if protocol_type:
            try:
                protocol_enum = ProtocolType(protocol_type)
                if protocol_enum in self.type_index:
                    type_docs = self.type_index[protocol_enum]
                    candidate_docs = candidate_docs.intersection(type_docs)
            except ValueError:
                logger.warning(f"Unknown protocol type: {protocol_type}")
        
        return candidate_docs
    
    def _calculate_relevance_score(self, query: str, search_terms: set, doc: ProtocolDocument) -> float:
        """Calcular score de relevancia para un documento."""
        score = 0.0
        
        # Texto completo para búsqueda
        full_text = f"{doc.title} {doc.content}".lower()
        query_lower = query.lower()
        
        # Score por coincidencia exacta en título (peso alto)
        if query_lower in doc.title.lower():
            score += 2.0
        
        # Score por términos en título
        title_words = set(doc.title.lower().split())
        title_matches = search_terms.intersection(title_words)
        score += len(title_matches) * 0.5
        
        # Score por términos en contenido
        content_words = set(full_text.split())
        content_matches = search_terms.intersection(content_words)
        score += len(content_matches) * 0.3
        
        # Score por densidad de términos
        if search_terms:
            density = len(content_matches) / len(search_terms)
            score += density * 0.4
        
        # Bonus por términos médicos específicos
        medical_bonus_terms = ['lpp', 'lesion', 'presion', 'tratamiento', 'protocolo']
        for term in medical_bonus_terms:
            if term in query_lower and term in full_text:
                score += 0.3
        
        # Normalizar score
        return min(score, 3.0) / 3.0
    
    def _find_matching_chunks(self, search_terms: set, chunks: List[str]) -> List[str]:
        """Encontrar chunks que coinciden con términos de búsqueda."""
        matching_chunks = []
        
        for chunk in chunks:
            chunk_lower = chunk.lower()
            chunk_words = set(chunk_lower.split())
            
            # Verificar si hay coincidencias
            matches = search_terms.intersection(chunk_words)
            if matches:
                matching_chunks.append((chunk, len(matches)))
        
        # Ordenar por número de coincidencias
        matching_chunks.sort(key=lambda x: x[1], reverse=True)
        
        return [chunk for chunk, _ in matching_chunks]
    
    def _create_content_snippet(self, query: str, content: str, max_length: int = 300) -> str:
        """Crear snippet del contenido centrado en la consulta."""
        query_terms = query.lower().split()
        content_lower = content.lower()
        
        # Buscar la primera ocurrencia de cualquier término de la consulta
        best_pos = -1
        for term in query_terms:
            pos = content_lower.find(term)
            if pos != -1:
                if best_pos == -1 or pos < best_pos:
                    best_pos = pos
        
        if best_pos == -1:
            # Si no se encuentra, usar el inicio
            snippet = content[:max_length]
        else:
            # Centrar alrededor de la posición encontrada
            start = max(0, best_pos - max_length // 2)
            end = start + max_length
            snippet = content[start:end]
            
            # Ajustar para no cortar palabras
            if start > 0:
                first_space = snippet.find(' ')
                if first_space != -1:
                    snippet = snippet[first_space + 1:]
            
            if end < len(content):
                last_space = snippet.rfind(' ')
                if last_space != -1:
                    snippet = snippet[:last_space]
        
        return snippet.strip()
    
    def _serialize_search_result(self, result: ProtocolSearchResult) -> Dict[str, Any]:
        """Serializar resultado de búsqueda para compatibilidad."""
        return {
            "doc_id": result.doc_id,
            "title": result.title,
            "content": result.content_snippet,
            "confidence": result.relevance_score,
            "sources": ["protocol_indexer"],
            "references": result.metadata.get("references", []),
            "protocol_type": result.protocol_type.value,
            "metadata": result.metadata,
            "matched_chunks": result.matched_chunks,
            "language": result.language,
            "last_updated": result.metadata.get("extraction_date")
        }
    
    def _get_search_cache_key(self, query: str, protocol_type: Optional[str], 
                            query_type: Optional[str], max_results: int, 
                            language: Optional[str]) -> str:
        """Generar clave de cache para búsqueda."""
        cache_data = f"{query}|{protocol_type or ''}|{query_type or ''}|{max_results}|{language or ''}"
        return hashlib.sha256(cache_data.encode()).hexdigest()[:16]
    
    async def _persist_protocol_document(self, doc: ProtocolDocument):
        """Persistir documento de protocolo en Redis."""
        if not self.redis_client:
            return
        
        try:
            doc_data = self._serialize_protocol_document(doc)
            key = f"protocol_doc:{doc.doc_id}"
            await self.redis_client.hset(key, mapping=doc_data)
            
            # Establecer TTL si se configura
            if hasattr(self.config, 'document_ttl') and self.config.document_ttl:
                await self.redis_client.expire(key, self.config.document_ttl)
            
        except Exception as e:
            logger.warning(f"Failed to persist protocol document {doc.doc_id}: {e}")
    
    def _serialize_protocol_document(self, doc: ProtocolDocument) -> Dict[str, str]:
        """Serializar documento de protocolo."""
        return {
            "doc_id": doc.doc_id,
            "title": doc.title,
            "content": doc.content,
            "protocol_type": doc.protocol_type.value,
            "metadata": json.dumps(doc.metadata),
            "chunks": json.dumps(doc.chunks),
            "language": doc.language,
            "file_path": doc.file_path or "",
            "created_at": doc.created_at.isoformat(),
            "last_updated": doc.last_updated.isoformat()
        }
    
    def _deserialize_protocol_document(self, doc_data: Dict[str, str]) -> Optional[ProtocolDocument]:
        """Deserializar documento de protocolo."""
        try:
            return ProtocolDocument(
                doc_id=doc_data["doc_id"],
                title=doc_data["title"],
                content=doc_data["content"],
                protocol_type=ProtocolType(doc_data["protocol_type"]),
                metadata=json.loads(doc_data["metadata"]),
                chunks=json.loads(doc_data["chunks"]),
                language=doc_data["language"],
                file_path=doc_data["file_path"] if doc_data["file_path"] else None,
                created_at=datetime.fromisoformat(doc_data["created_at"]),
                last_updated=datetime.fromisoformat(doc_data["last_updated"])
            )
        except Exception as e:
            logger.error(f"Failed to deserialize protocol document: {e}")
            return None
    
    async def index_sample_protocols(self):
        """Indexar protocolos de muestra para testing."""
        sample_protocols = [
            {
                "title": "Protocolo de Prevención de LPP en UCI",
                "content": """
                Protocolo para la prevención de lesiones por presión en unidad de cuidados intensivos.
                
                1. Evaluación de riesgo:
                - Aplicar escala de Braden cada 24 horas
                - Identificar factores de riesgo específicos
                - Documentar en historia clínica electrónica
                
                2. Medidas preventivas:
                - Cambios posturales cada 2 horas
                - Uso de superficies de redistribución de presión
                - Cuidados especializados de la piel
                - Nutrición adecuada e hidratación
                
                3. Monitoreo:
                - Inspección diaria de la piel
                - Documentación fotográfica si hay cambios
                - Reevaluación de riesgo semanal
                """,
                "protocol_type": ProtocolType.PREVENTION,
                "metadata": {
                    "department": "UCI",
                    "evidence_level": "high",
                    "last_review": "2024-01-01",
                    "references": ["NPUAP Guidelines 2019", "EPUAP Guidelines 2019"]
                }
            },
            {
                "title": "Tratamiento de LPP Grado 3 - Protocolo Avanzado",
                "content": """
                Protocolo para el tratamiento de lesiones por presión grado 3.
                
                Evaluación inicial:
                - Medición de dimensiones (largo x ancho x profundidad)
                - Evaluación del lecho de la herida
                - Identificación de tejidos comprometidos
                - Cultivo si signos de infección
                
                Manejo de la herida:
                - Desbridamiento quirúrgico si necesario
                - Limpieza con suero fisiológico a presión
                - Apósitos de espuma o alginato según exudado
                - Consideración de terapia de presión negativa (VAC)
                
                Manejo del dolor:
                - Analgesia previa a curaciones
                - Técnicas de distracción
                - Evaluación regular con escala EVA
                
                Soporte nutricional:
                - Evaluación nutricional completa
                - Suplementación proteica
                - Vitamina C y Zinc
                - Hidratación adecuada
                """,
                "protocol_type": ProtocolType.TREATMENT,
                "metadata": {
                    "complexity": "high",
                    "evidence_level": "high",
                    "specialties": ["wound_care", "surgery", "nutrition"],
                    "references": ["Cochrane Reviews 2023", "WOCN Guidelines"]
                }
            },
            {
                "title": "Protocolo de Antibióticos en LPP Infectadas",
                "content": """
                Protocolo para el manejo antimicrobiano de lesiones por presión infectadas.
                
                Criterios de infección:
                - Celulitis perilesional
                - Aumento del exudado purulento
                - Signos sistémicos de infección
                - Cultivo positivo
                
                Terapia empírica:
                - Amoxicilina-clavulánico 875/125mg cada 8 horas
                - Clindamicina 300-450mg cada 6 horas (alérgicos a penicilina)
                - Ceftriaxona 1-2g cada 24 horas (casos severos)
                
                Terapia dirigida:
                - S. aureus MSSA: Cefazolina 1g cada 8 horas
                - S. aureus MRSA: Vancomicina 15-20mg/kg cada 12 horas
                - P. aeruginosa: Piperacilina-tazobactam 4.5g cada 6 horas
                
                Duración: 7-14 días según respuesta clínica
                """,
                "protocol_type": ProtocolType.MEDICATION,
                "metadata": {
                    "specialty": "infectology",
                    "evidence_level": "high",
                    "drug_interactions": True,
                    "references": ["IDSA Guidelines 2024", "Antimicrobial Stewardship 2024"]
                }
            }
        ]
        
        indexed_count = 0
        for protocol in sample_protocols:
            try:
                await self.index_document_content(
                    content=protocol["content"],
                    title=protocol["title"],
                    protocol_type=protocol["protocol_type"],
                    metadata=protocol["metadata"]
                )
                indexed_count += 1
            except Exception as e:
                logger.error(f"Failed to index sample protocol: {e}")
        
        logger.info(f"Indexed {indexed_count} sample protocols")
        return indexed_count
    
    async def get_stats(self) -> Dict[str, Any]:
        """Obtener estadísticas del indexador."""
        return {
            "documents_indexed": self.stats["documents_indexed"],
            "total_chunks": self.stats["total_chunks"],
            "search_queries": self.stats["search_queries"],
            "cache_hits": self.stats["cache_hits"],
            "extraction_errors": self.stats["extraction_errors"],
            "cache_hit_rate": self.stats["cache_hits"] / max(self.stats["search_queries"], 1),
            "supported_formats": self.config.supported_formats,
            "redis_available": self.redis_client is not None,
            "protocol_types": {pt.value: len(docs) for pt, docs in self.type_index.items()}
        }
    
    async def cleanup(self):
        """Limpiar recursos."""
        try:
            if self.redis_client:
                await self.redis_client.close()
            
            # Limpiar caches
            self.search_cache.clear()
            self.term_index.clear()
            self.type_index.clear()
            self.protocol_store.clear()
            
            logger.info("Protocol indexer cleanup completed")
            
        except Exception as e:
            logger.error(f"Cleanup failed: {e}")


class ProtocolIndexerFactory:
    """Factory para crear instancias del indexador de protocolos."""
    
    @staticmethod
    async def create_indexer(config: Optional[ProtocolIndexConfig] = None) -> EnhancedProtocolIndexer:
        """Crear indexador enhanced."""
        indexer = EnhancedProtocolIndexer(config)
        await indexer.initialize()
        
        # Indexar protocolos de muestra si no hay documentos
        stats = await indexer.get_stats()
        if stats["documents_indexed"] == 0:
            await indexer.index_sample_protocols()
        
        return indexer